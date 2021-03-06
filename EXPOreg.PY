##I was tasked to create an registration and check-in application for our annual PNA Expo event.
##This was put on ice due to Covid
##I split the app into two parts, EXPOreg for geristering clients invted and RSVPd,
##the EXPOcheck part was for the check-in of the customers to avoid sharing of invitations to uninvited guestes.


import barcode
import os
from barcode.writer import ImageWriter
import sqlite3
from tkinter import *
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import SQLapp
import emailing

##f = open('barcodes - Copy.csv', 'r')
##contents = f.readlines()
##print(contents)

# i = contents[1]
# print(i)
# EAN = barcode.get_barcode_class('ean13')
# ean = EAN(str(i[:-2]), writer=ImageWriter())
# fullname = ean.save('barcode')

# x = 1
# for i in contents[1:]:
#     print(str(i[:-2]))
#     EAN = barcode.get_barcode_class('ean13')
#     ean = EAN(str(i[:-2]))
#     fullname = ean.save(str(i[:-2]))

titles = ['VIP BUYER', 'PNA BUYER', 'STAFF', 'EXIBITIONER', 'IT TECHNICIAN', 'TEST']


window=Tk()

window.configure(background = "white")

window.wm_title("EXPO")

##create document after registration
def create_regletter(bcode):
    EAN = barcode.get_barcode_class('ean13')
    ean = EAN(str(bcode), writer=ImageWriter())
    fullname = ean.save('./docs/'+str(bcode))
    document = Document('Expotemplate.docx')
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(fullname, width=Inches(3.0))
    # document.add_picture('barcode.png', width=Inches(3.0))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save('./docs/'+str(bcode)+'.docx')

##fill fields and drop down with selected details
def get_selected_row(event):
    if len(list1.curselection())> 0:
        fill_entries()
        global row_tuple
        row_tuple = list1.get(list1.curselection()[0])

##clear text fields
def clear_entries():
    # e1.delete(0, END)
    barcode_text.set("")
    e2.delete(0, END)
    # e3.delete(0, END)
    # e4.delete(0, END)
    checked_text.set("")
    e5.delete(0, END)

##fill text field with selected data
def fill_entries():
    # e1.delete(0, END)
    # e1.insert(END, list1.get(list1.curselection()[0])[1])
    barcode_text.set(list1.get(list1.curselection()[0])[1])
    e2.delete(0, END)
    e2.insert(END, list1.get(list1.curselection()[0])[2])
    # e3.delete(0, END)
    # e3.insert(END, list1.get(list1.curselection()[0])[4])
    title_text.set(list1.get(list1.curselection()[0])[4])
    # e4.delete(0, END)
    # e4.insert(END, list1.get(list1.curselection()[0])[5])
    checked_text.set(list1.get(list1.curselection()[0])[5])
    e5.delete(0, END)
    e5.insert(END, list1.get(list1.curselection()[0])[3])

##show all SQL data in display window
def show_command():
    list1.delete(0, END)
    for row in SQLapp.show_data():
        list1.insert(END, row)

##show searched SQL data in display window
def search_command():
    list1.delete(0, END)
    for row in SQLapp.search_data(barcode_text.get(), name_text.get().upper(), title_text.get().upper()):
        list1.insert(END, row)
    clear_entries()

##show searched SQL data and select first row on pressing 'enter'
def search_event(event):
    list1.delete(0, END)
    for row in SQLapp.search_data(barcode_text.get(), name_text.get().upper(), title_text.get().upper()):
        list1.insert(END, row)
    clear_entries()
    list1.select_set(0)
    fill_entries()

##create new registation in database
def insert_command():
    this_barcode = SQLapp.search_barcodedata()[0][0]
    # for code in SQLapp.search_barcodedata():
    #     barcodes_left.append(code[0])
    # print(barcode_text.get())
    # print(barcodes_left[0])
    list1.delete(0, END)
    if len(name_text.get()) > 0 and len(title_text.get()) > 0:
        SQLapp.delete_data(this_barcode)
##        check if customer is from a school or not and add as needed
        if len(school_text.get()) ==0:
            SQLapp.insert_data(this_barcode, name_text.get().upper(), "NONE", title_text.get().upper())
            create_regletter(this_barcode)
            emailing.send_email(str(email_address_text.get()), this_barcode)
            for row in SQLapp.search_data(this_barcode, name_text.get(), title_text.get()):
                list1.insert(END, row)
        else:
            SQLapp.insert_data(this_barcode, name_text.get().upper(), school_text.get(), title_text.get().upper())
            create_regletter(this_barcode)
            emailing.send_email(str(email_address_text.get()), this_barcode)
            for row in SQLapp.search_data(this_barcode, name_text.get(), title_text.get()):
                list1.insert(END, row)
    else:
        window2 = Tk()
        window2.wm_title('ERROR')
        E1 = Label(window2, text="Please check your info")
        E1.grid(row=0, column=0)
    clear_entries()

# def insert_command():
#     list1.delete(0, END)
#     create_regletter(barcode_text.get())
#     emailing.send_email(name_text.get(), barcode_text.get())
#     for row in SQLapp.search_data(barcode_text.get(), name_text.get(), title_text.get()):
#         list1.insert(END, row)
#     clear_entries()


##this part is for EXPOcheck part to mark client aa enetered or deny second entry
def update_command():
    if checked_text.get() == 'FALSE':
        SQLapp.update_data(row_tuple[0], barcode_text.get(), name_text.get(), school_text.get(), title_text.get())
        list1.delete(0, END)
        for row in SQLapp.search_data(barcode_text.get(), name_text.get(), title_text.get(), school_text.get(), checked_text.get()):
            list1.insert(END, row)
        # get_selected_row(self)
    else:
        window2 = Tk()
        window2.wm_title('ERROR')
        E1 = Label(window2, text="Client Already Entered")
        E1.grid(row=0, column=0)

##SQL delete for admin users only
def delete_command():
    SQLapp.delete_data(list1.get(list1.curselection()[0])[0])
    show_command()

##the rest is window layout and function links

l1 = Label(window, text="BARCODE", width=25)
l1.grid(row=0, column=0)

# l1.configure(background = "red")
l2 = Label(window, text="NAME", width=25)
l2.grid(row=1, column=0)
l3 = Label(window, text="TITLE", width=25)
l3.grid(row=3, column=0)
l5 = Label(window, text="School", width=25)
l5.grid(row=2, column=0)
l4 = Label(window, text="CHECKED", width=25)
l4.grid(row=0, column=2)
l6 = Label(window, text="Email", width=25)
l6.grid(row=1, column=2)

barcode_text=StringVar()
e1=Label(window, textvariable=barcode_text)
e1.grid(row=0, column=1)
name_text=StringVar()
e2=Entry(window, textvariable=name_text)
e2.grid(row=1, column=1)
school_text=StringVar()
e5=Entry(window, textvariable=school_text)
e5.grid(row=2, column=1)
title_text=StringVar(window)
title_text.set(titles[0])
# e3 = Entry(window, textvariable=title_text)
e3 = OptionMenu(window, title_text, *titles)
e3.grid(row=3, column=1)
checked_text=StringVar()
e4=Label(window, textvariable=checked_text)
e4.grid(row=0, column=3)
email_address_text = StringVar()
e6 = Entry(window, textvariable=email_address_text)
e6.grid(row=1, column=3)


list1 = Listbox(window, height=11, width=55)
list1.grid(row=5, column=2, columnspan=2, rowspan=3)

sb1 = Scrollbar(window)
sb1.grid(row=5, column=4, rowspan=3)

list1.configure(yscrollcommand=sb1.set)
sb1.configure(comman=list1.yview)
list1.bind('<<ListboxSelect>>', get_selected_row)


# b1=Button(window, text="Show All", command=show_command, width=16)
# b1.grid(row=5, column=0)
b2=Button(window, text="Search entry", command=search_command, width=16)
b2.grid(row=5, column=0)
window.bind('<Return>', search_event)
b3=Button(window, text="Add Entry", command=insert_command, width=16)
b3.grid(row=6, column=0)
b4=Button(window, text="Update", command=update_command, width=16)
b4.grid(row=7, column=0)
# b5=Button(window, text="Delete", command=delete_command, width=16)
# b5.grid(row=6, column=3)
# b6=Button(window, text="Close", command=window.destroy, width=16)
# b6.grid(row=7, column=3)
b5 = Button(window, text="Clear", command=clear_entries, width=16)
b5.grid(row=8, column=0)

window.mainloop()
