import barcode
import os
from barcode.writer import ImageWriter
import sqlite3
from tkinter import *
import SQLapp
import emailing
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

bcode = 1234567891234
# EAN = barcode.get_barcode_class('ean13')
# ean = EAN(str(bcode), writer=ImageWriter())
# fullname = ean.save('./docs/'+str(bcode))

document = Document('ExpoNameTag.docx')
p1 = document.add_paragraph()
r1= p1.add_run()
r1.font.size = Pt(26)
r1.font.bold = True
r1.add_text("Name")
p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2 = document.add_paragraph()
r2 = p2.add_run()
r2.font.size = Pt(22)
r2.add_text("School")
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3 = document.add_paragraph()
r3 = p3.add_run()
r3.font.size = Pt(26)
r3.font.bold = True
r3.add_text("Title")
p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# r = p.add_run()
# # r.add_picture('./docs/'+str(bcode)+'.png', width=Inches(3.0))
# # document.add_picture('barcode.png', width=Inches(3.0))
# r.add_text("Name")
# r.add_text("School")
# r.add_text("Title")
# p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.save('./docs/tag'+str(bcode)+'.docx')
if os.path.isfile('./docs/tag'+str(bcode)+'.docx') == True:
    os.startfile('./docs/tag'+str(bcode)+'.docx', "print")
else:
    time.sleep(1)
    continue
#
# emailing.send_email('orders@pnastrandgroup.co.za')



# x = []
# for code in SQLapp.search_barcodedata():
#     x.append(code[0])
# print(x)
# print('6103730387150' in str(x))

# os.startfile("365.html", "print")

# print(SQLapp.search_barcodedata()[0][0])
